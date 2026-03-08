from datetime import datetime
import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from drafter.document_types import DOCUMENT_TYPES

DOC_RGB_COLORS = {
    "FIR": (192, 57, 43),
    "Legal Notice": (41, 128, 185),
    "RTI Application": (39, 174, 96),
    "Consumer Complaint": (243, 156, 18),
    "Bail Application": (142, 68, 173),
}


def add_horizontal_rule(doc: Document):
    """Add a horizontal line to Word document."""
    paragraph = doc.add_paragraph()
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return paragraph


def generate_docx(
    document_type: str,
    document_text: str,
    extracted_fields: dict,
) -> bytes:
    """
    Generates a professional editable Word document.
    Returns bytes for Streamlit download_button.
    """
    doc = Document()
    doc_config = DOCUMENT_TYPES[document_type]
    rgb = DOC_RGB_COLORS.get(document_type, (51, 51, 51))
    doc_color = RGBColor(*rgb)
    today = datetime.today().strftime("%d %B %Y")

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    # Title
    title = doc.add_heading("", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"{doc_config['icon']} {document_type.upper()}")
    run.font.color.rgb = doc_color
    run.font.size = Pt(20)
    run.font.bold = True

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(doc_config["full_name"]).font.size = Pt(12)

    act_para = doc.add_paragraph()
    act_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    act_run = act_para.add_run(f"Under: {doc_config['act']}")
    act_run.font.size = Pt(9)
    act_run.font.italic = True

    conf_para = doc.add_paragraph()
    conf_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    conf_run = conf_para.add_run(
        "WITHOUT PREJUDICE | PRIVATE & CONFIDENTIAL"
    )
    conf_run.font.color.rgb = RGBColor(192, 57, 43)
    conf_run.font.size = Pt(8)
    conf_run.font.bold = True

    add_horizontal_rule(doc)

    # Metadata table
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.style = "Table Grid"
    meta_data = [
        ("Date", extracted_fields.get("date", today)),
        ("Document Type", doc_config["full_name"]),
        ("Authority", doc_config["authority"]),
        ("Urgency", extracted_fields.get("urgency_level", "MEDIUM")),
    ]
    for i, (label, value) in enumerate(meta_data):
        label_cell = meta_table.rows[i].cells[0]
        value_cell = meta_table.rows[i].cells[1]
        label_run = label_cell.paragraphs[0].add_run(label + ":")
        label_run.font.bold = True
        label_run.font.color.rgb = doc_color
        label_run.font.size = Pt(9)
        value_cell.paragraphs[0].add_run(value).font.size = Pt(9)

    doc.add_paragraph()

    # Body
    paragraphs = document_text.split("\n")
    for para_text in paragraphs:
        para_text = para_text.strip()
        if not para_text:
            doc.add_paragraph()
            continue

        is_header = (
            (para_text.isupper() and len(para_text) > 3)
            or (para_text.endswith(":") and len(para_text) < 60)
            or para_text.startswith("THAT")
            or para_text.startswith("GROUND")
            or para_text.startswith("PRAYER")
            or para_text.startswith("VERIFICATION")
        )

        if is_header:
            h = doc.add_heading(para_text, level=2)
            for run in h.runs:
                run.font.color.rgb = doc_color
                run.font.size = Pt(11)
        else:
            clean = para_text.replace("**", "").replace("__", "")
            p = doc.add_paragraph(clean)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in p.runs:
                run.font.size = Pt(10)

    # Applicable sections
    add_horizontal_rule(doc)
    sec_heading = doc.add_heading("APPLICABLE LEGAL SECTIONS", level=2)
    for run in sec_heading.runs:
        run.font.color.rgb = doc_color

    applicable = extracted_fields.get("applicable_sections", [])
    sec_text = (
        ", ".join(applicable)
        if isinstance(applicable, list)
        else str(applicable)
    )
    doc.add_paragraph(sec_text)

    # Disclaimer
    add_horizontal_rule(doc)
    disclaimer = doc.add_paragraph(
        "⚠️ DISCLAIMER: This document was AI-generated for reference only. "
        "Not legal advice. Verify with a qualified advocate before submission."
    )
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in disclaimer.runs:
        run.font.size = Pt(7)
        run.font.color.rgb = RGBColor(153, 153, 153)
        run.font.italic = True

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

